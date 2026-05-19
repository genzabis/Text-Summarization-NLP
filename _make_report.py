"""Generate Laporan_Project_IndoSum.docx dengan dokumentasi project lengkap."""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------- helpers ----------
BLUE = RGBColor(0x1F, 0x5A, 0xD6)
INK = RGBColor(0x0F, 0x17, 0x2A)
MUTE = RGBColor(0x64, 0x74, 0x8B)


def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = BLUE if level <= 2 else INK
    return h


def add_para(doc, text, bold=False, italic=False, size=11, color=INK,
             align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        b = p.add_run(bold_prefix)
        b.bold = True
        b.font.size = Pt(11)
        b.font.name = "Calibri"
        rest = p.add_run(text)
        rest.font.size = Pt(11)
        rest.font.name = "Calibri"
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.font.name = "Calibri"
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.name = "Calibri"
    return p


def add_code(doc, text):
    """Paragraf dengan font monospace + background abu-abu muda."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1F, 0x5A, 0xD6)
    # background
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F1F5FA")
    pPr.append(shd)
    return p


def add_table(doc, header, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # header
    for i, txt in enumerate(header):
        cell = table.cell(0, i)
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10.5)
        run.font.name = "Calibri"
        set_cell_bg(cell, "1F5AD6")

    # rows
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = table.cell(r, c)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10.5)
            run.font.name = "Calibri"
            run.font.color.rgb = INK

    # width
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = w
    return table


# =================================================================
def main():
    doc = Document()

    # ---------- page setup ----------
    section = doc.sections[0]
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # ---------- normal style ----------
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ============ COVER ============
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("LAPORAN PROYEK\nPemrosesan Bahasa Alami")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = BLUE
    r.font.name = "Calibri"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Aplikasi Peringkas Teks Bahasa Indonesia Berbasis IndoSum\n"
                    "TextRank · Seq2Seq + Bahdanau Attention · Google Gemini API")
    r.italic = True
    r.font.size = Pt(13)
    r.font.color.rgb = MUTE
    r.font.name = "Calibri"

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Kelompok 2")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = INK
    r.font.name = "Calibri"

    members = [
        "Atika Andrian Asmiran      — 234110601056",
        "Fatiyatul Amelia            — 234110601065",
        "Fikri Nofan Dwi Andika      — 234110601067",
        "Niamilah Nabil Syahputra    — 234110601087",
        "Novian Affan Ashofah        — 234110601088",
    ]
    for m in members:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(m)
        r.font.size = Pt(11.5)
        r.font.color.rgb = INK
        r.font.name = "Calibri"

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Mata Kuliah Pemrosesan Bahasa Alami\nSemester 6 · 2026")
    r.font.size = Pt(11)
    r.font.color.rgb = MUTE
    r.font.name = "Calibri"

    doc.add_page_break()

    # ============ 1. RINGKASAN EKSEKUTIF ============
    add_heading(doc, "1. Ringkasan Eksekutif", level=1)
    add_para(doc,
        "Proyek ini membangun sebuah aplikasi peringkas teks (text "
        "summarization) Bahasa Indonesia secara end-to-end. Aplikasi "
        "menggabungkan tiga pendekatan dalam satu antarmuka: TextRank "
        "sebagai baseline ekstraktif, model Seq2Seq + Bahdanau attention "
        "sebagai model utama yang dilatih sendiri di atas dataset "
        "IndoSum, serta integrasi Google Gemini API sebagai pembanding "
        "berbasis LLM. Sistem terdiri dari backend Flask (REST API), "
        "frontend HTML/CSS/JS murni dengan tema biru soft, serta pipeline "
        "training PyTorch yang reproducible.")

    add_para(doc,
        "Hasil evaluasi pada test set IndoSum (200 sampel test.01) "
        "menunjukkan TextRank mencapai ROUGE-1 = 0.4305, ROUGE-2 = 0.3005, "
        "dan ROUGE-L = 0.3542. Model Seq2Seq versi “tiny” (4.56 M "
        "parameter, dilatih 4 epoch pada CPU dengan 2.000 sampel) hanya "
        "mencapai ROUGE-1 ≈ 0.0011 — model masih under-fit. Skenario ini "
        "didokumentasikan secara jujur sebagai limitasi training di CPU "
        "dengan dataset dan epoch terbatas; pipeline-nya sendiri sudah "
        "siap untuk dilatih ulang di GPU/Colab dengan dataset penuh.")

    # ============ 2. LATAR BELAKANG ============
    add_heading(doc, "2. Latar Belakang & Motivasi", level=1)
    add_para(doc,
        "Volume teks digital berbahasa Indonesia tumbuh sangat cepat — "
        "berita, artikel, dan dokumen resmi menumpuk lebih cepat dari "
        "kemampuan manusia membacanya. Text summarization menjawab "
        "kebutuhan ini dengan menyaring poin-poin penting dari dokumen "
        "panjang menjadi ringkasan yang singkat namun tetap informatif. "
        "Pada konteks AI generatif, ringkasan yang baik juga "
        "memperkecil bias dan halusinasi: input yang lebih ringkas "
        "membantu model bahasa fokus pada inti, bukan pada detail "
        "yang tidak relevan.")

    add_para(doc,
        "Tugas ini menjadi tugas akhir mata kuliah Pemrosesan Bahasa "
        "Alami dengan fokus mengimplementasikan model Seq2Seq encoder-"
        "decoder serta membandingkannya dengan baseline TextRank dan "
        "LLM komersial (Gemini).")

    # ============ 3. TUJUAN ============
    add_heading(doc, "3. Tujuan Proyek", level=1)
    add_bullet(doc, "Membangun model peringkas teks Bahasa Indonesia berbasis Seq2Seq + Bahdanau attention.")
    add_bullet(doc, "Menjadikan TextRank sebagai baseline pembanding kuantitatif (ROUGE).")
    add_bullet(doc, "Mengintegrasikan Google Gemini API sebagai pembanding LLM abstraktif.")
    add_bullet(doc, "Menyediakan antarmuka web yang clean, ringan, dan mudah dipakai untuk demo.")
    add_bullet(doc, "Menyediakan pipeline training, evaluasi (ROUGE), dan benchmarking yang reproducible.")

    # ============ 4. DATASET ============
    add_heading(doc, "4. Dataset: IndoSum", level=1)
    add_para(doc,
        "IndoSum adalah dataset peringkasan berita berbahasa Indonesia "
        "yang dirilis melalui Kaggle (linkgish/indosum). Dataset ini "
        "berisi pasangan artikel berita lengkap dan ringkasan emas "
        "(gold summary) yang dibuat manusia. Format file adalah "
        ".jsonl, di mana setiap baris adalah objek JSON dengan struktur:")
    add_code(doc,
        '{\n'
        '  "id": "...",\n'
        '  "category": "...",\n'
        '  "source": "...",\n'
        '  "paragraphs": [[ [token, ...], ... ], ...],\n'
        '  "summary": [[token, ...], ...]\n'
        '}'
    )
    add_para(doc, "Statistik split yang dipakai pada proyek ini:")
    add_table(doc,
        ["Split", "File", "Jumlah Sampel", "Kegunaan"],
        [
            ["train", "train.01–05.jsonl", "± 14.000+", "Latih Seq2Seq"],
            ["dev",   "dev.01–05.jsonl",   "± 750",     "Validasi (early stopping)"],
            ["test",  "test.01–05.jsonl",  "± 3.762 (test.01)", "Evaluasi akhir / benchmark"],
        ])

    add_para(doc,
        "Pre-processing yang dilakukan: setiap artikel di-flatten dari "
        "struktur paragraph→sentence→token menjadi list kalimat, lalu "
        "tanda baca seperti spasi sebelum koma/titik dirapikan. "
        "Untuk Seq2Seq, kalimat-kalimat tersebut digabung menjadi "
        "single string sebelum di-tokenize ulang dengan vocabulary "
        "kustom (lihat Bagian 6).")

    # ============ 5. ARSITEKTUR SISTEM ============
    add_heading(doc, "5. Arsitektur Sistem", level=1)
    add_para(doc,
        "Sistem dibagi menjadi tiga lapis: backend, model layer, dan "
        "frontend. Komunikasi antar lapis menggunakan REST API berbasis "
        "JSON.")

    add_heading(doc, "5.1. Struktur Folder", level=2)
    add_code(doc,
        "indosum/\n"
        "├── train.0X.jsonl, dev.0X.jsonl, test.0X.jsonl   (dataset)\n"
        "└── app/\n"
        "    ├── backend/\n"
        "    │   ├── app.py                  (Flask + REST API)\n"
        "    │   ├── data_utils.py           (loader IndoSum)\n"
        "    │   ├── textrank_summarizer.py  (baseline TextRank)\n"
        "    │   ├── seq2seq_summarizer.py   (wrapper inferensi Seq2Seq)\n"
        "    │   ├── gemini_summarizer.py    (wrapper Gemini API)\n"
        "    │   └── rouge_eval.py           (ROUGE F1)\n"
        "    ├── models/\n"
        "    │   ├── seq2seq.py              (arsitektur GRU + attention)\n"
        "    │   ├── vocab.py                (tokenizer + Vocabulary)\n"
        "    │   └── checkpoints/            (seq2seq_best.pt, vocab.json)\n"
        "    ├── frontend/\n"
        "    │   ├── index.html, styles.css, app.js\n"
        "    ├── train.py                    (training Seq2Seq)\n"
        "    ├── benchmark.py                (ROUGE benchmark)\n"
        "    ├── requirements.txt\n"
        "    └── .env.example                (slot GEMINI_API_KEY)"
    )

    add_heading(doc, "5.2. Endpoint REST", level=2)
    add_table(doc,
        ["Method", "Endpoint", "Fungsi"],
        [
            ["GET",  "/api/status",
             "Status TextRank, Seq2Seq, dan Gemini (ready/off)."],
            ["POST", "/api/summarize",
             "Body { text, model, num_sentences, reference? } — meringkas + ROUGE opsional."],
            ["GET",  "/api/sample?split=&index=",
             "Mengambil sample IndoSum untuk demo."],
            ["POST", "/api/evaluate",
             "Hitung ROUGE antara prediksi dan referensi manual."],
        ],
        col_widths=[Cm(2.0), Cm(5.5), Cm(8.5)])

    # ============ 6. METODE & MODEL ============
    add_heading(doc, "6. Metode & Model", level=1)

    add_heading(doc, "6.1. Baseline · TextRank", level=2)
    add_para(doc,
        "TextRank adalah algoritma graf-based yang merepresentasikan "
        "tiap kalimat sebagai simpul, dan kemiripan TF-IDF + cosine "
        "antar kalimat sebagai bobot edge. Algoritma PageRank "
        "diiterasikan pada graf tersebut sampai konvergen, lalu "
        "kalimat-kalimat dengan skor tertinggi diambil sebagai "
        "ringkasan. Pendekatan ini bersifat ekstraktif (mengambil "
        "kalimat asli, tidak membentuk kalimat baru).")
    add_para(doc,
        "Implementasi pada proyek ini menggunakan numpy + tokenizer "
        "sederhana untuk Bahasa Indonesia. Tidak ada training; "
        "TextRank langsung jalan kapan pun.")

    add_heading(doc, "6.2. Model Utama · Seq2Seq + Bahdanau Attention", level=2)
    add_para(doc, "Arsitektur model utama:")
    add_bullet(doc,
        "encoder GRU bidirectional dengan padding mask — "
        "menerima token id dari artikel.",
        bold_prefix="Encoder · ")
    add_bullet(doc,
        "Bahdanau additive attention. Pada setiap langkah dekode, "
        "vektor query (hidden decoder) dipasangkan dengan setiap "
        "encoder hidden state untuk membentuk context vector dinamis.",
        bold_prefix="Attention · ")
    add_bullet(doc,
        "GRU unidirectional yang menerima [embedding token sebelumnya; "
        "context vector] dan menghasilkan distribusi probabilitas "
        "token berikutnya melalui linear layer + softmax.",
        bold_prefix="Decoder · ")
    add_bullet(doc,
        "vocab kustom dari korpus IndoSum, special tokens "
        "<pad>, <bos>, <eos>, <unk>; tokenizer sederhana berbasis "
        "regex (lower-case, pisah tanda baca).",
        bold_prefix="Tokenizer · ")
    add_bullet(doc,
        "teacher forcing 0.85 pada training, greedy decoding pada "
        "inferensi (max-tgt-len = 25 token untuk versi tiny).",
        bold_prefix="Strategi · ")

    add_heading(doc, "6.2.1. Memahami Bahdanau Attention", level=2)
    add_para(doc,
        "Bahdanau attention adalah mekanisme attention pertama untuk "
        "model Seq2Seq, diperkenalkan oleh Dzmitry Bahdanau, Kyunghyun "
        "Cho, dan Yoshua Bengio pada paper “Neural Machine Translation "
        "by Jointly Learning to Align and Translate” (2014). Mekanisme "
        "ini menjadi cikal bakal seluruh keluarga model attention modern, "
        "termasuk Transformer.")

    add_para(doc,
        "Masalah yang diselesaikan. Pada Seq2Seq vanilla, seluruh "
        "informasi artikel harus dipadatkan oleh encoder ke dalam satu "
        "vektor hidden state terakhir. Untuk artikel berita yang "
        "panjang, vektor ini menjadi bottleneck: informasi di kalimat "
        "awal cenderung hilang saat encoder mencapai kalimat akhir, "
        "sehingga ringkasan yang dihasilkan dekoder bias terhadap "
        "bagian akhir teks.",
        bold=False)

    add_para(doc,
        "Solusi. Bahdanau membiarkan dekoder “menengok ulang” semua "
        "hidden state encoder setiap kali akan menghasilkan satu kata. "
        "Dekoder belajar memberi bobot lebih besar pada bagian artikel "
        "yang paling relevan dengan kata yang sedang dihasilkan, dan "
        "membentuk context vector dinamis dari kombinasi tertimbang "
        "tersebut. Inilah yang disebut “alignment” antara token "
        "output dan token input.")

    add_heading(doc, "Cara Kerja Tahap demi Tahap", level=3)
    add_number(doc,
        "Encoder bidirectional GRU membaca artikel → menghasilkan satu "
        "hidden state untuk setiap token input: h₁, h₂, …, hₙ. Pada "
        "encoder bidirectional, setiap hᵢ menggabungkan konteks dari "
        "arah kiri-ke-kanan dan kanan-ke-kiri.")
    add_number(doc,
        "Saat dekoder akan menghasilkan token ke-t, dekoder punya "
        "hidden state s_{t-1} (state sebelum langkah t).")
    add_number(doc,
        "Untuk setiap hᵢ encoder, hitung skor kecocokan e_{t,i} "
        "antara s_{t-1} dan hᵢ menggunakan formula additive "
        "(ciri khas Bahdanau):")
    add_code(doc,
        "e_{t,i} = vᵀ · tanh( W₁ · s_{t-1} + W₂ · h_i )")
    add_number(doc,
        "Normalisasi skor dengan softmax → bobot attention "
        "α_{t,i}, dengan Σᵢ α_{t,i} = 1. Bobot ini bisa "
        "diinterpretasi sebagai “seberapa penting kata ke-i artikel "
        "untuk kata ke-t ringkasan”.")
    add_code(doc,
        "α_{t,i} = exp(e_{t,i}) / Σⱼ exp(e_{t,j})")
    add_number(doc,
        "Bentuk context vector dinamis c_t sebagai jumlah "
        "tertimbang dari semua encoder hidden state:")
    add_code(doc,
        "c_t = Σᵢ α_{t,i} · h_i")
    add_number(doc,
        "Dekoder menerima [embedding kata sebelumnya y_{t-1}; c_t] "
        "sebagai input GRU pada langkah t, lalu memproyeksikan ke "
        "distribusi probabilitas seluruh vocabulary lewat linear + "
        "softmax untuk memilih kata berikutnya y_t.")

    add_heading(doc, "Analogi Sederhana", level=3)
    add_para(doc,
        "Bayangkan seorang penerjemah manusia yang menerjemahkan "
        "artikel panjang ke ringkasan. Dia tidak menulis ringkasan "
        "“buta” cuma dari ingatan kabur tentang artikel — saat menulis "
        "kata ke-5 di ringkasan, dia melirik balik ke kalimat asli yang "
        "berkaitan dengan posisi tersebut. Bobot α_{t,i} adalah "
        "“lirikan” itu: kata input yang paling relevan diberi atensi "
        "tinggi, sisanya diabaikan.")

    add_heading(doc, "Bahdanau vs Luong Attention", level=3)
    add_para(doc,
        "Luong attention (2015) adalah variasi populer yang muncul "
        "satu tahun setelah Bahdanau. Perbedaannya:")
    add_table(doc,
        ["Aspek", "Bahdanau (Additive)", "Luong (Multiplicative)"],
        [
            ["Skor",
             "vᵀ tanh(W₁s + W₂h)",
             "sᵀ W h  (dot product berbobot)"],
            ["Hidden state yang dipakai",
             "s_{t-1} (state sebelum langkah t)",
             "s_t (state setelah langkah t)"],
            ["Komputasi",
             "Sedikit lebih lambat (tanh + 2 matmul + 1 vT)",
             "Lebih cepat (1 matmul saja)"],
            ["Performa",
             "Sangat baik untuk sequence panjang",
             "Mirip Bahdanau, kadang sedikit lebih baik di translation"],
            ["Implementasi proyek",
             "✓ digunakan",
             "—"],
        ],
        col_widths=[Cm(4.0), Cm(6.0), Cm(6.0)])

    add_heading(doc, "Mengapa Memilih Bahdanau di Proyek Ini?", level=3)
    add_bullet(doc,
        "menjadi standar “textbook” saat membahas "
        "attention-based Seq2Seq, sehingga pas untuk demo materi "
        "perkuliahan.",
        bold_prefix="Pedagogis · ")
    add_bullet(doc,
        "tahan untuk sequence panjang (artikel berita 100+ token), "
        "lebih stabil daripada Seq2Seq tanpa attention.",
        bold_prefix="Robust · ")
    add_bullet(doc,
        "bobot α_{t,i} bisa divisualisasikan sebagai matriks "
        "alignment artikel ↔ ringkasan — berguna untuk debugging.",
        bold_prefix="Interpretable · ")
    add_para(doc,
        "Pada implementasi PyTorch proyek ini, Bahdanau attention "
        "berada di file app/models/seq2seq.py kelas BahdanauAttention, "
        "dipanggil setiap langkah dekoder di kelas Seq2Seq.forward / "
        ".generate.")

    add_heading(doc, "6.3. Pendekatan LLM · Google Gemini API", level=2)
    add_para(doc,
        "Sebagai pembanding ketiga, aplikasi mengintegrasikan model "
        "gemini-1.5-flash dari Google. Prompt berisi instruksi "
        "Bahasa Indonesia: “Ringkas artikel berikut dalam N kalimat …”. "
        "API key dibaca dari file app/.env (variabel "
        "GEMINI_API_KEY); jika kosong, endpoint /api/summarize "
        "dengan model=gemini akan mengembalikan HTTP 503 dengan pesan "
        "yang informatif tanpa membuat sistem crash.")

    # ============ 7. TRAINING ============
    add_heading(doc, "7. Training Model Seq2Seq", level=1)

    add_heading(doc, "7.1. Konfigurasi Training", level=2)
    add_table(doc,
        ["Parameter", "Nilai", "Catatan"],
        [
            ["Sampel training",    "2.000",    "subset dari train.01.jsonl (CPU-friendly)"],
            ["Sampel validasi",    "100",      "subset dari dev.01.jsonl"],
            ["Vocabulary size",    "10.000",   "min frekuensi token disesuaikan otomatis"],
            ["Embedding dimension","48",       "tiny config"],
            ["Hidden dimension",   "96",       "tiny config"],
            ["Max source length",  "120 token","artikel dipotong/pad"],
            ["Max target length",  "25 token", "ringkasan greedy"],
            ["Batch size",         "32",       ""],
            ["Optimizer",          "Adam",     "lr = 1e-3"],
            ["Teacher forcing",    "0.85",     "decay opsional"],
            ["Loss",               "CrossEntropy + ignore_index <pad>", ""],
            ["Epoch",              "4",        "early stop manual via val_loss"],
            ["Device",             "CPU",      "tidak ada GPU di mesin developer"],
            ["Total parameter",    "4.56 M",   ""],
        ],
        col_widths=[Cm(4.5), Cm(4.0), Cm(7.5)])

    add_heading(doc, "7.2. Log Training", level=2)
    add_para(doc, "Log per epoch (loss CE, durasi pada CPU):")
    add_code(doc,
        "[info] device = cpu\n"
        "[info] train pairs: 2000\n"
        "[info] dev pairs: 100\n"
        "[info] vocab size: 10000\n"
        "[info] model params: 4.56 M\n"
        "[epoch 1/4] train_loss=7.6903  val_loss=6.9395  time= 95.6s  -> save best\n"
        "[epoch 2/4] train_loss=6.9253  val_loss=6.8022  time= 97.5s  -> save best\n"
        "[epoch 3/4] train_loss=6.7366  val_loss=6.7027  time=102.9s  -> save best\n"
        "[epoch 4/4] train_loss=6.5580  val_loss=6.6276  time=106.4s  -> save best\n"
        "[done] best val_loss = 6.6276")

    add_para(doc,
        "Total waktu training: ± 6 menit 43 detik untuk 4 epoch pada "
        "CPU. Loss validasi turun monoton 6.94 → 6.80 → 6.70 → 6.63, "
        "menunjukkan model masih bisa belajar lebih jauh — namun "
        "untuk konfigurasi tiny ini kecepatan konvergensi terbatas "
        "oleh kapasitas model + ukuran data.")

    add_heading(doc, "7.3. Catatan Training Sebelumnya", level=2)
    add_para(doc,
        "Versi pertama model dilatih dengan emb_dim=128, hidden_dim=256, "
        "max_src=400, max_tgt=80, dataset 4.000 sampel × 3 epoch. "
        "Pada CPU developer, satu epoch saja membutuhkan > 30 menit "
        "tanpa progress yang terlihat dan training dihentikan. "
        "Versi tiny kemudian dipilih demi reproducibility di mesin "
        "tanpa GPU.")

    # ============ 8. EVALUASI ============
    add_heading(doc, "8. Evaluasi & Hasil", level=1)
    add_para(doc,
        "Evaluasi menggunakan metrik ROUGE-1, ROUGE-2, dan ROUGE-L "
        "(F-measure) dari library rouge-score. Skor dihitung pada "
        "subset test.01.jsonl dataset IndoSum.")

    add_heading(doc, "8.1. Hasil Kuantitatif", level=2)
    add_table(doc,
        ["Model", "Sampel", "ROUGE-1", "ROUGE-2", "ROUGE-L"],
        [
            ["TextRank (baseline)",      "200", "0.4305", "0.3005", "0.3542"],
            ["Seq2Seq tiny (4 epoch CPU)","50",  "0.0011", "0.0000", "0.0011"],
            ["Google Gemini 1.5 Flash",  "—",   "*tidak diukur otomatis*",
             "*tidak diukur otomatis*", "*tidak diukur otomatis*"],
        ],
        col_widths=[Cm(5.5), Cm(2.0), Cm(2.6), Cm(2.6), Cm(2.6)])

    add_para(doc,
        "Catatan: Gemini tidak diikutkan dalam benchmarking otomatis "
        "karena API berbayar dan harus menggunakan API key user. "
        "Evaluasi Gemini bisa dilakukan manual lewat tombol "
        "“Hitung ROUGE vs reference” di antarmuka.")

    add_heading(doc, "8.2. Analisis Kualitatif", level=2)
    add_para(doc,
        "Pada artikel test ke-0 (kasus berita tentang “wanita terberat "
        "di dunia”), TextRank mampu mengangkat 3 kalimat paling "
        "representatif dengan tepat (gold summary mencantumkan kasus "
        "Eman Ahmed Abd El Aty di Mumbai dan TextRank pun memilih "
        "kalimat-kalimat yang sama). Seq2Seq versi tiny, sebaliknya, "
        "menghasilkan rentetan token <unk> karena vocabulary terbatas "
        "10k dan dataset training yang sangat kecil — model belum "
        "punya distribusi token yang mature.")

    add_heading(doc, "8.3. Mengapa Seq2Seq Skor-nya Rendah?", level=2)
    add_bullet(doc,
        "training di CPU sangat membatasi ukuran "
        "dataset dan model. Bahdanau attention adalah O(T²) terhadap "
        "panjang sequence dan tidak teroptimasi tanpa GPU.",
        bold_prefix="Compute · ")
    add_bullet(doc,
        "2.000 sampel × 25 target token × 4 epoch tidak cukup untuk "
        "model menghafal struktur kalimat ringkasan Indonesia.",
        bold_prefix="Volume data · ")
    add_bullet(doc,
        "vocab 10k tidak menutupi sebagian "
        "besar token nama diri / kata jarang dalam berita; banyak "
        "token jadi <unk>.",
        bold_prefix="Vocabulary · ")
    add_bullet(doc,
        "max_tgt = 25 token bias menjadikan "
        "model belajar stop dini, sehingga output sering pendek "
        "dan didominasi <unk>.",
        bold_prefix="Sequence length · ")
    add_para(doc,
        "Solusi yang sudah disiapkan dalam pipeline: argumen "
        "--resume di train.py, sehingga training tinggal dilanjutkan "
        "di Google Colab dengan GPU dan dataset penuh tanpa perlu "
        "rebuild vocabulary.")

    # ============ 9. UI ============
    add_heading(doc, "9. Antarmuka Pengguna", level=1)
    add_para(doc,
        "Frontend dibangun dengan HTML/CSS/JS murni (tanpa framework) "
        "menggunakan tema light dengan single accent biru (#2f6feb). "
        "Memenuhi spesifikasi minimal yang diberikan:")

    add_table(doc,
        ["Spesifikasi", "Implementasi pada UI"],
        [
            ["Judul aplikasi",
             "Topbar tag “Tugas Pemrosesan Bahasa Alami · Kelompok 2” + heading utama."],
            ["Deskripsi singkat",
             "Subtitle menjelaskan dataset IndoSum dan tiga pendekatan model."],
            ["Input area",
             "Textarea besar + tombol Unggah .txt + char counter realtime."],
            ["Tombol Proses",
             "Tombol primer “Proses Ringkasan” dengan icon ▶."],
            ["Area output",
             "Card terpisah “Hasil Ringkasan” + tombol Salin yang muncul setelah hasil."],
            ["Confidence/score",
             "Tiga kartu ROUGE-1/-2/-L + hint penjelasan skala 0–1."],
            ["Logo & identitas kelompok",
             "Logo SVG dokumen+ceklis biru di sidebar; identitas kelompok di tab Tentang."],
        ],
        col_widths=[Cm(4.5), Cm(11.5)])

    add_para(doc,
        "Selain tab Ringkas Teks, terdapat tab Sample Dataset (telusuri "
        "shard IndoSum langsung dari .jsonl) dan tab Tentang. Sidebar "
        "menampilkan status real-time ketiga model (TextRank ready / "
        "Seq2Seq ready / Gemini off bila API key kosong).")

    # ============ 10. CARA MENJALANKAN ============
    add_heading(doc, "10. Cara Menjalankan", level=1)
    add_number(doc, "Install dependency: pip install -r app/requirements.txt")
    add_number(doc, "(Opsional) salin app/.env.example menjadi app/.env, isi GEMINI_API_KEY untuk mengaktifkan Gemini.")
    add_number(doc, "(Opsional) latih ulang Seq2Seq:")
    add_code(doc,
        "python -m app.train --indosum-dir . \\\n"
        "    --epochs 4 --max-train 2000 --max-val 100 \\\n"
        "    --batch-size 32 --emb-dim 48 --hidden-dim 96 \\\n"
        "    --max-src-len 120 --max-tgt-len 25 \\\n"
        "    --vocab-size 10000 --lr 1e-3 --teacher-forcing 0.85")
    add_number(doc, "Jalankan server: python -m app.backend.app")
    add_number(doc, "Buka browser di http://127.0.0.1:5000")
    add_number(doc, "Reproduksi benchmark:")
    add_code(doc,
        "python -m app.benchmark --indosum-dir . --max 200 --models textrank\n"
        "python -m app.benchmark --indosum-dir . --max 50 --models seq2seq \\\n"
        "    --checkpoint app/models/checkpoints/seq2seq_best.pt \\\n"
        "    --vocab app/models/checkpoints/vocab.json")

    # ============ 11. KESIMPULAN ============
    add_heading(doc, "11. Kesimpulan & Saran", level=1)
    add_para(doc,
        "Aplikasi peringkas teks Bahasa Indonesia berhasil dibangun "
        "secara end-to-end (model + backend + frontend) dengan tiga "
        "pendekatan: TextRank, Seq2Seq + Bahdanau attention, dan "
        "Google Gemini API. Pipeline training Seq2Seq sudah terbukti "
        "konvergen (val_loss turun monoton selama 4 epoch), namun "
        "kualitas output masih rendah karena keterbatasan training "
        "di CPU dengan data dan epoch yang kecil. Baseline TextRank "
        "tetap menjadi pilihan praktis terbaik di lingkungan tanpa "
        "GPU dengan ROUGE-1 sekitar 0.43 — angka yang konsisten "
        "dengan literatur peringkasan ekstraktif untuk Bahasa "
        "Indonesia.")

    add_para(doc, "Saran pengembangan lanjutan:")
    add_bullet(doc, "Latih ulang Seq2Seq di GPU/Colab dengan ≥ 30.000 sampel × ≥ 10 epoch, emb=128, hidden=256.")
    add_bullet(doc, "Eksplorasi model transformer (mBART/T5) yang sudah pre-trained Bahasa Indonesia.")
    add_bullet(doc, "Tambah beam search (beam=4) untuk dekoder Seq2Seq.")
    add_bullet(doc, "Implementasi cakupan ROUGE per-kategori berita (politik, olahraga, dll).")
    add_bullet(doc, "Tambah unit-test untuk fungsi pre-processing dan ROUGE.")

    # ============ 12. REFERENSI ============
    add_heading(doc, "12. Referensi Singkat", level=1)
    add_bullet(doc, "Mihalcea & Tarau (2004) — TextRank: Bringing Order into Texts.")
    add_bullet(doc, "Bahdanau, Cho, Bengio (2014) — Neural Machine Translation by Jointly Learning to Align and Translate.")
    add_bullet(doc, "Sutskever, Vinyals, Le (2014) — Sequence to Sequence Learning with Neural Networks.")
    add_bullet(doc, "Lin (2004) — ROUGE: A Package for Automatic Evaluation of Summaries.")
    add_bullet(doc, "Kurniawan & Louvan (2018) — IndoSum: A New Benchmark Dataset for Indonesian Text Summarization.")
    add_bullet(doc, "Google AI (2024) — gemini-1.5-flash via google-generativeai SDK.")

    # ---------- save ----------
    out = "Laporan_Project_IndoSum.docx"
    doc.save(out)
    print(f"[done] saved -> {out}")


if __name__ == "__main__":
    main()
